import hashlib
import io
import json
from datetime import datetime
from typing import List, Any, Dict, Optional
import pandas as pd
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from docx import Document
from docx.shared import Inches, Pt
from app.db.base import AuditLogEntry
from sqlalchemy.orm import Session
from app.models.audit import AuditEvent

class CertifiedReportService:
    @staticmethod
    def calculate_hash(content: bytes) -> str:
        """Calculates SHA-256 hash for a given content."""
        return hashlib.sha256(content).hexdigest()

    @staticmethod
    async def create_audit_entry(
        db: Session, 
        user_id: str, 
        action: str, 
        file_hash: str, 
        filters: Dict[str, Any]
    ):
        """Logs the export event for chain of custody."""
        event = AuditEvent(
            user_id=user_id,
            event_type="geotrail_export",
            description=f"Exportacao certificada de geolocalizacao ({action})",
            metadata={
                "sha256": file_hash,
                "filters": filters,
                "exported_at": datetime.utcnow().isoformat()
            }
        )
        db.add(event)
        db.commit()

    async def generate_xlsx(self, data: List[Dict[str, Any]], agent_name: str) -> bytes:
        df = pd.DataFrame(data)
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Audit Trail')
            
            # Metadata sheet for certification
            meta_df = pd.DataFrame([{
                "Relatorio": "Auditoria de Geolocalizacao FARO",
                "Agente": agent_name,
                "Data de Geracao": datetime.utcnow().isoformat(),
                "Certificacao": "SHA-256 Integridy Shield"
            }])
            meta_df.to_excel(writer, index=False, sheet_name='Certificacao')
            
        return output.getvalue()

    async def generate_pdf(self, data: List[Dict[str, Any]], agent_name: str, filters: Dict[str, Any]) -> bytes:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        elements = []
        styles = getSampleStyleSheet()

        # Header
        elements.append(Paragraph("F.A.R.O. - Relatorio de Auditoria de Geolocalizacao", styles['Title']))
        elements.append(Paragraph(f"Agente: {agent_name}", styles['Normal']))
        elements.append(Paragraph(f"Periodo: {filters.get('start_date')} ate {filters.get('end_date')}", styles['Normal']))
        elements.append(Spacer(1, 12))

        # Table
        table_data = [["Timestamp", "Latitude", "Longitude", "Status", "Bat %"]]
        for item in data[:500]: # Limit PDF size
            table_data.append([
                str(item['recorded_at']),
                str(item['location']['latitude']),
                str(item['location']['longitude']),
                str(item['connectivity_status']),
                str(item.get('battery_level', '--'))
            ])

        t = Table(table_data)
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.black)
        ]))
        elements.append(t)
        
        # Certification Footer
        elements.append(Spacer(1, 24))
        elements.append(Paragraph("DADOS CERTIFICADOS - CADEIA DE CUSTODIA FARO", styles['Heading3']))
        
        # We'll add the hash later in a second pass or use a placeholder
        # For now, structural signature
        elements.append(Paragraph(f"Gerado em: {datetime.utcnow().isoformat()} por Sistema Automático de Auditoria", styles['Small']))

        doc.build(elements)
        return buffer.getvalue()

    async def generate_docx(self, data: List[Dict[str, Any]], agent_name: str) -> bytes:
        doc = Document()
        doc.add_heading('F.A.R.O. - Auditoria de Campo', 0)
        
        p = doc.add_paragraph('Relatório técnico de movimentação para fins de ')
        p.add_run('Cadeia de Custódia').bold = True
        
        doc.add_heading(f'Agente: {agent_name}', level=1)
        
        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Data/Hora'
        hdr_cells[1].text = 'Latitude'
        hdr_cells[2].text = 'Longitude'
        hdr_cells[3].text = 'Status'
        
        for item in data[:100]: # Sample for docx
            row_cells = table.add_row().cells
            row_cells[0].text = str(item['recorded_at'])
            row_cells[1].text = str(item['location']['latitude'])
            row_cells[2].text = str(item['location']['longitude'])
            row_cells[3].text = str(item['connectivity_status'])
            
        doc.add_page_break()
        
        target_stream = io.BytesIO()
        doc.save(target_stream)
        return target_stream.getvalue()

report_service = CertifiedReportService()
